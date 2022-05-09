import io
from docx import Document
from django import template
from django.shortcuts import render, redirect, get_object_or_404
from .forms import BuiltMaterialForm, BuiltUEForm, ExcavationForm, FactForm, \
                   InclusionForm, RoomForm, SedimentaryMaterialForm, SedimentaryUEForm, \
                   PhotoForm, CustomUserCreationForm, CustomUserChangeForm
from .models import Excavation, Photo, Fact, Inclusion, Room, BuiltMaterial, \
                    SedimentaryMaterial, BuiltUE, SedimentaryUE
from django.conf import settings
from django.template.loader import get_template
from django.contrib.auth.forms import PasswordResetForm
from django.contrib.auth.models import User
from django.contrib.auth.tokens import default_token_generator
from django.utils.http import urlsafe_base64_encode
from django.utils.encoding import force_bytes
from django.contrib.auth.decorators import login_required, permission_required, user_passes_test
from django.core.mail import BadHeaderError, send_mail
from django.http import HttpResponse
from django.core.exceptions import PermissionDenied
from rest_framework import viewsets
from rest_framework.authtoken.views import ObtainAuthToken
from rest_framework.authtoken.models import Token
from rest_framework.response import Response
from .serializers import ExcavationSerializer, PhotoSerializer,SedimentaryMaterialSerializer,\
                         BuiltMaterialSerializer, SedimentaryUESerializer, BuiltUESerializer,\
                         InclusionSerializer, RoomSerializer, SedimentaryUEFloorSerializer, \
                         SedimentaryUESectionSerializer, BuiltUEFloorSerializer, \
                         BuiltUESectionSerializer, FactSerializer, FactPlanSerializer, \
                         FactSectionSerializer, RoomFloorSerializer

# Create your views here.
def index(request):
    return render(request, 'home.html')

def about(request):
    return render(request, 'about.html')

def contact(request):
    return render(request, 'contact.html')

def team(request):
    return render(request, 'team.html')

# ############
# EXCAVATIONS
# ############
@login_required
@permission_required('myFindings.view_excavation', raise_exception=True)
def list_allexcavations(request):
    # Get all excavations
    excavations = Excavation.objects.all()
    data = { 'excavations': excavations}

    return render(request, 'excavations_list.html', data)

@login_required
@permission_required('myFindings.add_excavation', raise_exception=True)
def add_excavation(request):
    # Get the fields of excavation
    data = { 'form': ExcavationForm() }

    if request.method == 'POST':
        # Get the data entered by the user
        form = ExcavationForm(data=request.POST)
        if(form.is_valid()):    # Check if valid
            form.save()         # Save form

            # Redirect to the list of excavations
            return redirect(to='excavations')
        else:
            data['form'] = form

    return render(request, 'add_excavation.html', data)

@login_required
@permission_required('myFindings.change_excavation', raise_exception=True)
def modify_excavation(request, id):
    excavation = get_object_or_404(Excavation, id=id)
    
    # Guardar el formulario con los datos del cuadro
    data = { 'form': ExcavationForm(instance=excavation) }

    if request.method == 'POST':
        form = ExcavationForm(data=request.POST, instance=excavation)
        if form.is_valid():       # Si es válido
            form.save()           # Guardarlo

            return redirect(to="excavations")

        data["form"] = form

    return render(request, 'modify_excavation.html', data)

@login_required
@permission_required('myFindings.delete_excavation', raise_exception=True)
def delete_excavation(request, id):
    # Get the excavation, if it doesn't exist, get an Http404
    excavation = get_object_or_404(Excavation, id=id)

    # Delete the excavation
    excavation.delete()    

    return redirect(to="excavations")    

# ################
# SEDIMENTARY UE
# ################
@login_required
@permission_required('myFindings.view_sedimentaryue', raise_exception=True)
def list_allsedimentaryues(request):
    # Get all sedimentary ues
    sedimentaryues = SedimentaryUE.objects.all()
    data = { 'sedimentaryues': sedimentaryues}

    return render(request, 'sedimentaryues_list.html', data)

@login_required
@permission_required('myFindings.add_sedimentaryue', raise_exception=True)
def add_sedimentaryue(request):
    # Get the fields of the sedimentary ue
    data = { 'form': SedimentaryUEForm() }

    if request.method == 'POST':
        # Get the data entered by the user
        form = SedimentaryUEForm(data=request.POST, files=request.FILES)
        if(form.is_valid()):    # Check if valid
            form.save()         # Save form

            # Redirect to the list of excavations
            return redirect(to='sedimentaryues')
        else:
            data['form'] = form

    return render(request, 'add_sedimentaryue.html', data)

@login_required
@permission_required('myFindings.change_sedimentaryue', raise_exception=True)
def modify_sedimentaryue(request, id):
    sedimentaryue = get_object_or_404(SedimentaryUE, id=id)
    
    # Guardar el formulario con los datos del cuadro
    data = { 'form': SedimentaryUEForm(instance=sedimentaryue) }

    if request.method == 'POST':
        form = SedimentaryUEForm(data=request.POST, instance=sedimentaryue, files=request.FILES)
        if form.is_valid():       # Si es válido
            form.save()           # Guardarlo

            return redirect(to="sedimentaryues")

        data["form"] = form

    return render(request, 'modify_sedimentaryue.html', data)

@login_required
@permission_required('myFindings.delete_sedimentaryue', raise_exception=True)
def delete_sedimentaryue(request, id):
    # Get the sedimentary ue, if it doesn't exist, get an Http404
    sedimentaryue = get_object_or_404(SedimentaryUE, id=id)

    # Delete the sedimentaryue
    sedimentaryue.delete()    

    return redirect(to="sedimentaryues")   

# ################
# BUILT UE
# ################
@login_required
@permission_required('myFindings.view_builtue', raise_exception=True)
def list_allbuiltues(request):
    # Get all built ues
    builtues = BuiltUE.objects.all()
    data = { 'builtues': builtues}

    return render(request, 'builtues_list.html', data)

@login_required
@permission_required('myFindings.add_builtue', raise_exception=True)
def add_builtue(request):
    # Get the fields of built ue
    data = { 'form': BuiltUEForm() }

    if request.method == 'POST':
        # Get the data entered by the user
        form = BuiltUEForm(data=request.POST, files=request.FILES)
        if(form.is_valid()):    # Check if valid
            form.save()         # Save form

            # Redirect to the list of excavations
            return redirect(to='builtues')
        else:
            data['form'] = form

    return render(request, 'add_builtue.html', data)

@login_required
@permission_required('myFindings.change_builtue', raise_exception=True)
def modify_builtue(request, id):
    builtue = get_object_or_404(BuiltUE, id=id)
    
    # Guardar el formulario con los datos del cuadro
    data = { 'form': BuiltUEForm(instance=builtue) }

    if request.method == 'POST':
        form = BuiltUEForm(data=request.POST, instance=builtue, files=request.FILES)
        if form.is_valid():       # Si es válido
            form.save()           # Guardarlo

            return redirect(to="builtues")

        data["form"] = form

    return render(request, 'modify_builtue.html', data)

@login_required
@permission_required('myFindings.delete_builtue', raise_exception=True)
def delete_builtue(request, id):
    # Get the sedimentary ue, if it doesn't exist, get an Http404
    builtue = get_object_or_404(BuiltUE, id=id)

    # Delete the excavation
    builtue.delete()    

    return redirect(to="builtues")


# ################
# FACT
# ################
@login_required
@permission_required('myFindings.view_fact', raise_exception=True)
def list_allfacts(request):
    # Get all facts
    facts = Fact.objects.all()
    data = { 'facts': facts}

    return render(request, 'facts_list.html', data)

@login_required
@permission_required('myFindings.add_fact', raise_exception=True)
def add_fact(request):
    # Get the fields of fact
    data = { 'form': FactForm() }

    if request.method == 'POST':
        # Get the data entered by the user
        form = FactForm(data=request.POST, files=request.FILES)
        if(form.is_valid()):    # Check if valid
            form.save()         # Save form

            # Redirect to the list of facts
            return redirect(to='facts')
        else:
            data['form'] = form

    return render(request, 'add_fact.html', data)

@login_required
@permission_required('myFindings.change_fact', raise_exception=True)
def modify_fact(request, id):
    fact = get_object_or_404(Fact, id=id)
    
    # Guardar el formulario con los datos del cuadro
    data = { 'form': FactForm(instance=fact) }

    if request.method == 'POST':
        form = FactForm(data=request.POST, instance=fact, files=request.FILES)
        if form.is_valid():       # Si es válido
            form.save()           # Guardarlo

            return redirect(to="facts")

        data["form"] = form

    return render(request, 'modify_fact.html', data)

@login_required
@permission_required('myFindings.delete_fact', raise_exception=True)
def delete_fact(request, id):
    # Get the fact, if it doesn't exist, get an Http404
    fact = get_object_or_404(Fact, id=id)

    # Delete the excavation
    fact.delete()    

    return redirect(to="facts")  

# ################
# ROOM
# ################
@login_required
@permission_required('myFindings.view_room', raise_exception=True)
def list_allrooms(request):
    # Get all rooms
    rooms = Room.objects.all()
    data = { 'rooms': rooms}

    return render(request, 'rooms_list.html', data)

@login_required
@permission_required('myFindings.add_room', raise_exception=True)
def add_room(request):
    # Get the fields of fact
    data = { 'form': RoomForm() }

    if request.method == 'POST':
        # Get the data entered by the user
        form = RoomForm(data=request.POST, files=request.FILES)
        if(form.is_valid()):    # Check if valid
            form.save()         # Save form

            # Redirect to the list of facts
            return redirect(to='rooms')
        else:
            data['form'] = form

    return render(request, 'add_room.html', data) 

@login_required
@permission_required('myFindings.change_room', raise_exception=True)
def modify_room(request, id):
    room = get_object_or_404(Room, id=id)
    
    # Guardar el formulario con los datos del cuadro
    data = { 'form': RoomForm(instance=room) }

    if request.method == 'POST':
        form = RoomForm(data=request.POST, instance=room, files=request.FILES)
        if form.is_valid():       # Si es válido
            form.save()           # Guardarlo

            return redirect(to="rooms")

        data["form"] = form

    return render(request, 'modify_room.html', data)

@login_required
@permission_required('myFindings.delete_room', raise_exception=True)
def delete_room(request, id):
    # Get the room, if it doesn't exist, get an Http404
    room = get_object_or_404(Room, id=id)

    # Delete the room
    room.delete()    

    return redirect(to="rooms") 

# ################
# PHOTO
# ################
@login_required
@permission_required('myFindings.view_photo', raise_exception=True)
def list_allphotos(request):
    # Get all photos
    photos = Photo.objects.all()
    data = { 'photos': photos}

    return render(request, 'photos_list.html', data)

@login_required
@permission_required('myFindings.add_photo', raise_exception=True)
def add_photo(request):
    # Get the fields of photo
    data = { 'form': PhotoForm() }

    if request.method == 'POST':
        # Get the data entered by the user
        form = PhotoForm(data=request.POST, files=request.FILES)
        if(form.is_valid()):    # Check if valid
            form.save()         # Save form

            # Redirect to the list of facts
            return redirect(to='photos')
        else:
            data['form'] = form

    return render(request, 'add_photo.html', data)

@login_required
@permission_required('myFindings.change_fotografia', raise_exception=True)
def modify_photo(request, id):
    photo = get_object_or_404(Photo, id=id)
    
    # Guardar el formulario con los datos del cuadro
    data = { 'form': PhotoForm(instance=photo) }

    if request.method == 'POST':
        form = PhotoForm(data=request.POST, instance=photo, files=request.FILES)
        if form.is_valid():       # Si es válido
            form.save()           # Guardarlo

            return redirect(to="photos")

        data["form"] = form

    return render(request, 'modify_photo.html', data)
 
@login_required
@permission_required('myFindings.delete_photo', raise_exception=True)
def delete_photo(request, id):
    # Get the photo, if it doesn't exist, get an Http404
    photo = get_object_or_404(Photo, id=id)

    # Delete the photo
    photo.delete()    

    return redirect(to="photos") 

# ################
# INCLUSION
# ################
@login_required
@permission_required('myFindings.view_inclusion', raise_exception=True)
def list_allinclusions(request):
    # Get all photos
    inclusions = Inclusion.objects.all()
    data = { 'inclusions': inclusions}

    return render(request, 'inclusions_list.html', data)

@login_required
@permission_required('myFindings.add_inclusion', raise_exception=True)
def add_inclusion(request):
    # Get the fields of inclusion
    data = { 'form': InclusionForm() }

    if request.method == 'POST':
        # Get the data entered by the user
        form = InclusionForm(data=request.POST, files=request.FILES)
        if(form.is_valid()):    # Check if valid
            form.save()         # Save form

            # Redirect to the list of facts
            return redirect(to='inclusions')
        else:
            data['form'] = form

    return render(request, 'add_inclusion.html', data)

@login_required
@permission_required('myFindings.change_inclusion', raise_exception=True)
def modify_inclusion(request, id):
    inclusion = get_object_or_404(Inclusion, id=id)
    
    # Guardar el formulario con los datos del cuadro
    data = { 'form': InclusionForm(instance=inclusion) }

    if request.method == 'POST':
        form = InclusionForm(data=request.POST, instance=inclusion, files=request.FILES)
        if form.is_valid():       # Si es válido
            form.save()           # Guardarlo

            return redirect(to="inclusions")

        data["form"] = form

    return render(request, 'modify_inclusion.html', data)

@login_required
@permission_required('myFindings.delete_inclusion', raise_exception=True)
def delete_inclusion(request, id):
    # Get the inclusion, if it doesn't exist, get an Http404
    inclusion = get_object_or_404(Inclusion, id=id)

    # Delete the inclusion
    inclusion.delete()    

    return redirect(to="inclusions")

# #####################
# SEDIMENTARY MATERIAL
# #####################
@login_required
@permission_required('myFindings.view_sedimentarymaterial', raise_exception=True)
def list_allsedimentarymaterials(request):
    # Get all sedimentary materials
    sedimentarymaterials = SedimentaryMaterial.objects.all()
    data = { 'sedimentarymaterials': sedimentarymaterials}

    return render(request, 'sedimentarymaterials_list.html', data)

@login_required
@permission_required('myFindings.add_sedimentarymaterial', raise_exception=True)
def add_sedimentarymaterial(request):
    # Get the fields of sedimentary materials
    data = { 'form': SedimentaryMaterialForm() }

    if request.method == 'POST':
        # Get the data entered by the user
        form = SedimentaryMaterialForm(data=request.POST)
        if(form.is_valid()):    # Check if valid
            form.save()         # Save form

            # Redirect to the list of facts
            return redirect(to='sedimentarymaterials')
        else:
            data['form'] = form

    return render(request, 'add_sedimentarymaterial.html', data)

@login_required
@permission_required('myFindings.delete_sedimentarymaterial', raise_exception=True)
def delete_sedimentarymaterial(request, nombre):
    # Get the sedimentary material, if it doesn't exist, get an Http404
    sedimentarymaterial = get_object_or_404(SedimentaryMaterial, nombre=nombre)

    # Delete the sedimentary material
    sedimentarymaterial.delete()    

    return redirect(to="sedimentarymaterials")

# #####################
# BUILT MATERIAL
# #####################
@login_required
@permission_required('myFindings.view_builtmaterial', raise_exception=True)
def list_allbuiltmaterials(request):
    # Get all built materials
    builtmaterials = BuiltMaterial.objects.all()
    data = { 'builtmaterials': builtmaterials}

    return render(request, 'builtmaterials_list.html', data)

@login_required
@permission_required('myFindings.add_builtmaterial', raise_exception=True)
def add_builtmaterial(request):
    # Get the fields of sedimentary materials
    data = { 'form': BuiltMaterialForm() }

    if request.method == 'POST':
        # Get the data entered by the user
        form = BuiltMaterialForm(data=request.POST)
        if(form.is_valid()):    # Check if valid
            form.save()         # Save form

            # Redirect to the list of facts
            return redirect(to='builtmaterials')
        else:
            data['form'] = form

    return render(request, 'add_builtmaterial.html', data)

@login_required
@permission_required('myFindings.delete_builtmaterial', raise_exception=True)
def delete_builtmaterial(request, nombre):
    # Get the built material, if it doesn't exist, get an Http404
    builtmaterial = get_object_or_404(BuiltMaterial, nombre=nombre)

    # Delete the built material
    builtmaterial.delete()    

    return redirect(to="builtmaterials")

# ####################
# SPECIFIC LISTINGS  
# ####################
@login_required
@permission_required('myFindings.view_excavation', raise_exception=True)
@permission_required('myFindings.view_sedimentaryue', raise_exception=True)
@permission_required('myFindings.view_builtue', raise_exception=True)
def list_excavationues(request, id):
    # Get the excavation
    excavation = get_object_or_404(Excavation, id=id)

    # Find all associated sedimentary stratigraphic units
    sedimentaryues = SedimentaryUE.objects.filter(excavacion__id=id)

    # Find all associated built stratigraphic units
    builtues = BuiltUE.objects.filter(excavacion__id=id)

    data = { 'sedimentaryues': sedimentaryues, 'builtues': builtues, 'n_excavacion': excavation.n_excavacion}

    return render(request, 'excavationues.html', data)

@login_required
@permission_required('myFindings.view_room', raise_exception=True)
@permission_required('myFindings.view_fact', raise_exception=True)
def list_roomfacts(request, id):
    # Get the room
    room = get_object_or_404(Room, id=id)

    # Find all associated facts
    facts = Fact.objects.filter(estancia__id=id)

    data = { 'facts': facts,  'n_room': room.n_estancia}

    return render(request, 'facts_list.html', data)

@login_required
@permission_required('myFindings.view_fact', raise_exception=True)
@permission_required('myFindings.view_sedimentaryue', raise_exception=True)
@permission_required('myFindings.view_builtue', raise_exception=True)
def list_factues(request, id):
    # Get the fact
    fact = get_object_or_404(Fact, id=id)

    # Find all associated sedimentary stratigraphic units
    sedimentaryues = SedimentaryUE.objects.filter(hecho__id=id)

    # Find all associated built stratigraphic units
    builtues = BuiltUE.objects.filter(hecho__id=id)

    nombre = fact.letra + fact.numero
    data = { 'sedimentaryues': sedimentaryues, 'builtues': builtues, 'n_hecho': nombre}

    return render(request, 'excavationues.html', data)

# ######################
# Django authentication
# ######################
def register(request):
    # Get the fields of the registration form
    data = { 'form': CustomUserCreationForm() }

    if request.method == 'POST':
        # Get the data entered by the user
        form = CustomUserCreationForm(data=request.POST)
        if(form.is_valid()):    # Check if valid
            user = form.save()         # Save form
            # default to non-active
            user.is_active = False
            user.save()

            # Send an email to the archeologist administering the site
            template_data = {
                'user': user,
                'protocol': 'http',
                'domain': request.get_host(),
            }
            send_email(subject='Verificación de registro', from_email=settings.EMAIL_HOST_USER,
                        recipient_list=[settings.EMAIL_HOST_USER], html_message='registration/registration_check.html', data=template_data)

            return render(request, 'registration/register_confirm.html')

        else:
            data['form'] = form

    return render(request, 'registration/register.html', data)

def send_email(subject, from_email, recipient_list, message='',
              fail_silently=False, auth_user=None, auth_password=None,
              connection=None, html_message=None, data=None):
              
    if (html_message is not None):
        template = get_template(html_message)

        if data is not None:
            html_message = template.render(data)
    
    # Send the email
    try:
        send_mail(subject, message, from_email, recipient_list,
                fail_silently=fail_silently, auth_user=auth_user,
                auth_password=auth_password, connection=connection,
                html_message=html_message)
    except BadHeaderError:
        return HttpResponse('Invalid header found.')

def send_email_password_reset(request):
    try:
        # Get the user
        user = User.objects.get(email=request.POST['email']) 
    except User.DoesNotExist:
        return redirect(to='password_reset_done')
    
    # Get the fields of the reset form
    data = { 'user': user,
             'protocol': 'http', 
             'domain': request.get_host(),
             'uid': urlsafe_base64_encode(force_bytes(user.pk)),
             'token': default_token_generator.make_token(user),
    }
    
    if request.method == 'POST':
        # Get the data entered by the user
        form = PasswordResetForm(data=request.POST)
        if(form.is_valid()):    # Check if valid
            send_email(subject='MyFindings: restablecer contraseña', from_email=settings.EMAIL_HOST_USER,
                        recipient_list=[user.email], html_message='registration/password_reset_email.html', data=data)

            return redirect(to='password_reset_done')
        else:
            data['form'] = form

    return redirect(to='password_reset', context=data)

# ####################
# AUXILIARY DECORATOR 
# ####################
def group_required(*group_names):
    def in_groups(u):
        if u.is_authenticated:
            if bool(u.groups.filter(name__in=group_names)) | u.is_superuser:
                return True
        # Raise a permission denied error instead of returning False
        raise PermissionDenied

    return user_passes_test(in_groups)

@login_required
@group_required('Staff')
def staff_panel(request):
    # Get all excavations
    users = User.objects.filter(is_superuser=False)
    staff_users = User.objects.filter(groups__name__in=['Staff'])
  
    # Exclude the users that are already staff
    users = users.exclude(id__in=staff_users.values_list('id', flat=True))
            
    data = { 'users': users}
    return render(request, 'staff_panel.html', data)

@login_required
@group_required('Staff')
def change_perms(request, id):
    # Get the user
    user = get_object_or_404(User, id=id)

    # Save the inicial value of is_active field
    was_active = user.is_active

    # Fill the form with the user's data
    data = { 'form': CustomUserChangeForm(instance=user) }

    # Get the fields for the emails
    template_data = {
        'user': user,
        'protocol': 'http',
        'domain': request.get_host(),
    }

    if request.method == 'POST':
        # Get the data entered by the user
        form = CustomUserChangeForm(data=request.POST, instance=user)
        
        if(form.is_valid()):        # Check if valid
            print(was_active)
            print(form.cleaned_data['is_active'])
            if was_active and not form.cleaned_data['is_active']:
                # If the user is active and is deactivated, send an email to the user              
                send_email(subject='MyFindings: cuenta desactivada', from_email=settings.EMAIL_HOST_USER,
                        recipient_list=[user.email], html_message='registration/account_deactivated.html', data=template_data)
            
            form.save()      # Save form

            # Check if the is_active field was changed
            if(form.cleaned_data['is_active']):

                # Send an email to the user
                send_email(subject='MyFindings: cuenta activada', from_email=settings.EMAIL_HOST_USER,
                        recipient_list=[user.email], html_message='registration/account_activated.html', data=template_data)
     
            return redirect(to='staff_panel')
        
        data['form'] = form

    return render(request, 'change_perms.html', data)


# ######################
# REPORT GENERATOR
# ######################
def generate_report(request, id):
    excavation = get_object_or_404(Excavation, id=id)   # Get the excavation   

    # Source: https://python-docx.readthedocs.io/en/latest/ - Temporary
    document = Document()                               # Create a new document

    document.add_heading('Informe de la excavación ' + str(excavation.n_excavacion), 0)           # Add a heading
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    # Save document to memory and download to the user's browser 
    document_data = io.BytesIO()
    document.save(document_data)
    document_data.seek(0)
    response = HttpResponse(document_data.getvalue())
    response["Content-Disposition"] = 'attachment; filename = "Informe de excavación.docx"'
    response["Content-Encoding"] = "UTF-8"
    
    return response

# ######################
# API REST
# ######################
class ExcavationViewSet(viewsets.ModelViewSet):
    queryset = Excavation.objects.all()
    serializer_class = ExcavationSerializer

class PhotoViewSet(viewsets.ModelViewSet):
    queryset = Photo.objects.all()
    serializer_class = PhotoSerializer

class InclusionViewSet(viewsets.ModelViewSet):
    queryset = Inclusion.objects.all()
    serializer_class = InclusionSerializer

class SedimentaryUEViewSet(viewsets.ModelViewSet):
    queryset = SedimentaryUE.objects.all()
    lookup_field = 'codigo'

    def get_serializer_class(self):
        if self.request.GET.get('type') == 'floor':
            return SedimentaryUEFloorSerializer
        elif(self.request.GET.get('type') == 'section'):
            return SedimentaryUESectionSerializer
        return SedimentaryUESerializer

class BuiltUEViewSet(viewsets.ModelViewSet):
    queryset = BuiltUE.objects.all()
    lookup_field = 'codigo'

    def get_serializer_class(self):
        if self.request.GET.get('type') == 'floor':
            return BuiltUEFloorSerializer
        elif(self.request.GET.get('type') == 'section'):
            return BuiltUESectionSerializer
        return BuiltUESerializer

class SedimentaryMaterialViewSet(viewsets.ModelViewSet):
    queryset = SedimentaryMaterial.objects.all()
    serializer_class = SedimentaryMaterialSerializer

class BuiltMaterialViewSet(viewsets.ModelViewSet):
    queryset = BuiltMaterial.objects.all()
    serializer_class = BuiltMaterialSerializer
    lookup_field = 'codigo'

class FactViewSet(viewsets.ModelViewSet):
    queryset = Fact.objects.all()
    serializer_class = FactSerializer

    def get_serializer_class(self):
        if self.request.GET.get('type') == 'plan':
            return FactPlanSerializer
        elif(self.request.GET.get('type') == 'section'):
            return FactSectionSerializer
        return FactSerializer

class RoomViewSet(viewsets.ModelViewSet):
    queryset = Room.objects.all()
    serializer_class = RoomSerializer

    def get_serializer_class(self):
        if self.request.GET.get('type') == 'floor':
            return RoomFloorSerializer
        return RoomSerializer

class CustomAuthToken(ObtainAuthToken):

    def post(self, request, *args, **kwargs):
        serializer = self.serializer_class(data=request.data,
                                           context={'request': request})
        serializer.is_valid(raise_exception=True)
        user = serializer.validated_data['user']
        token, created = Token.objects.get_or_create(user=user)
        return Response({
            'token': token.key,
            'user_id': user.pk,
            'username': user.username,
            'first_name': user.first_name,
            'last_name': user.last_name,
            'email': user.email,
            'groups': [group.name for group in user.groups.all()],
            'user_permissions': [permission for permission in user.get_all_permissions()]
        })